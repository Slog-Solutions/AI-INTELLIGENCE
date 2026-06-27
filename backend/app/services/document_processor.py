"""
ATIP Enterprise Document Processor
====================================
Production-grade document processing with:
- Heading-aware, paragraph-aware, table-aware, sentence-overlap chunking
- Military entity extraction (officers, ranks, battalions, units, ops, weapons, vehicles, locations, dates)
- Alias normalization (Captain Jatin Kumar Verma = Capt Verma = Officer Verma)
- Large PDF support (300+ pages per document, 5 simultaneous documents)
- Per-page and per-section metadata tracking for accurate citations
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import uuid4

import pandas as pd
import fitz  # PyMuPDF

from docx import Document as DocxDocument
from ..config import UPLOAD_DIR

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# Military Entity Patterns
# ─────────────────────────────────────────────
RANK_PATTERNS = [
    r'\b(General|Gen|Lieutenant General|Lt Gen|Major General|Maj Gen|Brigadier|Brig)',
    r'\b(Colonel|Col|Lieutenant Colonel|Lt Col|Major|Maj)',
    r'\b(Captain|Capt|Lieutenant|Lt|Second Lieutenant|2Lt|2nd Lt)',
    r'\b(Subedar Major|Sub Maj|Subedar|Sub|Naib Subedar|Nb Sub)',
    r'\b(Havildar Major|Hav Maj|Havildar|Hav|Naik|Nk|Lance Naik|L Nk)',
    r'\b(Sepoy|Sep|Rifleman|Rfn|Signalman|Sigmn|Sapper|Spr|Gunner|Gnr)',
]

UNIT_PATTERNS = [
    r'\b(\d+(?:st|nd|rd|th)?\s+(?:Infantry|Armoured|Artillery|Engineers?|Signals?|Corps?|Division|Brigade|Battalion|Regiment|Company|Platoon|Section|Squad))\b',
    r'\b([A-Z]{2,6}\s+(?:Battalion|Regiment|Brigade|Division|Corps|Company))\b',
    r'\b((?:Alpha|Bravo|Charlie|Delta|Echo|Foxtrot|Golf|Hotel|India|Juliet|Kilo|Lima|Mike|November|Oscar|Papa|Quebec|Romeo|Sierra|Tango|Uniform|Victor|Whiskey|Xray|Yankee|Zulu)\s+(?:Company|Platoon|Section|Team))\b',
]

OPERATION_PATTERNS = [
    r'\b(Op(?:eration)?(?:s)?\s+[A-Z][A-Za-z0-9\-]+)\b',
    r'\b(Exercise\s+[A-Z][A-Za-z0-9\-]+)\b',
]

WEAPON_PATTERNS = [
    r'\b(AK-?47|AK-?74|M4|M16|INSAS|Tavor|Sig Sauer|Glock|Beretta)\b',
    r'\b(RPG|ATGM|Stinger|Igla|Javelin|Carl Gustav|NLAW|Panzerfaust)\b',
    r'\b(Artillery|Howitzer|Mortar|Cannon|Machine [Gg]un|Sniper|Rifle|Pistol|Shotgun)\b',
    r'\b([A-Z]{2,5}-\d{1,4}[A-Z]?\s+(?:missile|rocket|round|shell|grenade|mine))\b',
]

VEHICLE_PATTERNS = [
    r'\b(T-?(?:54|55|62|64|72|80|90)|Arjun|Leclerc|Abrams|Leopard|Challenger|Merkava)\b',
    r'\b(BMP|Bradley|Warrior|Stryker|M113|BTR|MT-LB)\b',
    r'\b(APC|IFV|MBT|AFV|MRAP|JLTV|HMMWV|Humvee)\b',
    r'\b(Mi-?\d{2,3}|UH-?60|CH-?47|AH-?64|Dhruv|Rudra|Prachand)\b',
    r'\b((?:main battle|infantry fighting|armoured personnel)\s+(?:tank|carrier|vehicle))\b',
]

LOCATION_PATTERNS = [
    r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3}\s+(?:Valley|Ridge|Hill|Pass|Sector|Zone|Area|Post|Forward|Base|Camp|FOB|COP|OP))\b',
    r'\b(Grid\s+[A-Z]{1,2}\s*\d{4,8})\b',
    r'\b((?:North|South|East|West|NE|NW|SE|SW)\s+(?:Sector|Zone|Flank|Axis))\b',
]

DATE_PATTERNS = [
    r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
    r'\b(\d{1,2}\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{2,4})\b',
    r'\b((?:Q[1-4]|FY)\s*\d{2,4})\b',
]

# ─────────────────────────────────────────────
# Alias Normalization
# ─────────────────────────────────────────────
def _build_alias_map(entities: Dict[str, List[str]]) -> Dict[str, str]:
    """Build a canonical → alias map for person name normalization."""
    alias_map = {}
    officers = entities.get("officers", [])
    for name in officers:
        # Normalise: strip rank prefix and whitespace
        parts = name.strip().split()
        # Collect all sub-name combos as aliases of the full name
        if len(parts) >= 2:
            for i in range(1, len(parts)):
                alias = " ".join(parts[i:])
                if alias and alias not in alias_map:
                    alias_map[alias.lower()] = name
            # First name + last name
            alias_map[parts[-1].lower()] = name
    return alias_map


def _extract_officer_names(text: str) -> List[str]:
    """Extract person names preceded by military ranks."""
    names = []
    # Match Rank + Name(s)  e.g. "Captain Jatin Kumar Verma"
    pattern = re.compile(
        r'\b(?:General|Gen|Lieutenant General|Lt Gen|Major General|Maj Gen|Brigadier|Brig|'
        r'Colonel|Col|Lieutenant Colonel|Lt Col|Major|Maj|Captain|Capt|Lieutenant|Lt|'
        r'2nd Lt|Second Lieutenant|Subedar|Sub|Havildar|Hav|Naik|Nk|Lance Naik|Sepoy|Sep|'
        r'Rifleman|Rfn|Sapper|Spr|Gunner|Gnr)\s+'
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,4})',
        re.MULTILINE
    )
    for m in pattern.finditer(text):
        full = m.group(0).strip()
        if full not in names:
            names.append(full)
    return names


# ─────────────────────────────────────────────
# Heading Detection
# ─────────────────────────────────────────────
_HEADING_RE = re.compile(
    r'^(?:'
    r'\d+[\.\d]*\s+[A-Z]|'          # Numbered: 1. Introduction
    r'[IVX]+\.\s+[A-Z]|'            # Roman: III. Mission
    r'(?:SECTION|CHAPTER|ANNEX|APPENDIX|PART)\s+[A-Z0-9]|'
    r'[A-Z][A-Z\s]{4,}$'            # ALL CAPS heading
    r')',
    re.MULTILINE
)


def _is_heading(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 120:
        return False
    return bool(_HEADING_RE.match(line))


def _detect_table_block(lines: List[str], start_idx: int) -> Tuple[bool, int]:
    """Detect pipe-delimited or aligned-column table blocks."""
    if start_idx >= len(lines):
        return False, start_idx
    line = lines[start_idx]
    if '|' in line or re.match(r'^[\s\-\+]{5,}$', line):
        end = start_idx
        while end < len(lines) and ('|' in lines[end] or re.match(r'^[\s\-\+]+$', lines[end])):
            end += 1
        return True, end
    return False, start_idx


# ─────────────────────────────────────────────
# Main Processor Class
# ─────────────────────────────────────────────
class DocumentProcessor:

    # ── File I/O ──────────────────────────────
    @staticmethod
    def save_upload(file_obj, filename: str) -> str:
        safe_name = Path(filename).name
        target_path = Path(UPLOAD_DIR) / safe_name
        if target_path.exists():
            stem = target_path.stem
            suffix = target_path.suffix
            target_path = target_path.with_name(f"{stem}-{uuid4().hex[:8]}{suffix}")
        file_obj.seek(0)
        with open(target_path, "wb") as f:
            f.write(file_obj.read())
        return str(target_path)

    @staticmethod
    def extract_metadata(
        path: str,
        original_filename: str,
        category: str,
        source: str,
        page_count: int = 0,
        chunk_count: int = 0,
        entities: Optional[Dict] = None,
    ) -> Dict[str, Any]:
        file_path = Path(path)
        meta = {
            "original_filename": Path(original_filename).name,
            "saved_filename": file_path.name,
            "extension": file_path.suffix.lower(),
            "category": category,
            "source": source,
            "size_bytes": file_path.stat().st_size,
            "page_count": page_count,
            "chunk_count": chunk_count,
        }
        if entities:
            meta["entities"] = entities
        return meta

    @staticmethod
    def generate_preview(parsed: Dict[str, Any]) -> Dict[str, Any]:
        if parsed["type"] == "excel":
            preview = {sheet: data[:10] for sheet, data in parsed["sheets"].items()}
            return {"type": "excel", "data": preview}
        elif parsed["type"] == "csv":
            return {"type": "csv", "data": parsed["rows"][:10]}
        elif parsed["type"] == "pdf":
            return {"type": "pdf", "data": parsed["full_text"][:1000]}
        elif parsed["type"] == "docx":
            return {"type": "docx", "data": "\n".join(parsed["paragraphs"][:10])}
        return {"type": "unknown", "data": ""}

    # ── Parsers ──────────────────────────────
    @staticmethod
    def parse_excel(path: str) -> Dict[str, Any]:
        workbook = pd.ExcelFile(path)
        sheets = {}
        for sheet in workbook.sheet_names:
            df = workbook.parse(sheet_name=sheet)
            sheets[sheet] = df.fillna("").to_dict(orient="records")
        return {"type": "excel", "sheets": sheets}

    @staticmethod
    def parse_csv(path: str) -> Dict[str, Any]:
        df = pd.read_csv(path)
        return {"type": "csv", "rows": df.fillna("").to_dict(orient="records")}

    @staticmethod
    def parse_pdf(path: str) -> Dict[str, Any]:
        """
        Extract text from PDF page-by-page, preserving page numbers.
        Handles 300+ page documents without memory issues.
        """
        page_texts: List[Dict[str, Any]] = []
        full_text_parts: List[str] = []
        page_count = 0
        try:
            with fitz.open(path) as doc:
                page_count = doc.page_count
                for page_num in range(page_count):
                    page = doc.load_page(page_num)
                    # Extract blocks with positions for layout-aware chunking
                    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                    raw_text = page.get_text("text")
                    
                    # Collect headings detected on this page
                    headings_on_page: List[str] = []
                    for block in blocks.get("blocks", []):
                        if block.get("type") == 0:  # text block
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    span_text = span.get("text", "").strip()
                                    span_size = span.get("size", 0)
                                    span_flags = span.get("flags", 0)
                                    is_bold = bool(span_flags & 2**4)
                                    # Large or bold text → likely heading
                                    if span_size >= 13 or is_bold:
                                        if span_text and len(span_text) <= 120:
                                            headings_on_page.append(span_text)
                    
                    page_texts.append({
                        "page_number": page_num + 1,
                        "text": raw_text,
                        "headings": headings_on_page,
                    })
                    full_text_parts.append(raw_text)

            return {
                "type": "pdf",
                "full_text": "\n".join(full_text_parts),
                "page_texts": page_texts,
                "page_count": page_count,
            }
        except Exception as e:
            logger.error(f"PDF parse error for {path}: {e}")
            return {"type": "pdf", "full_text": "", "page_texts": [], "page_count": 0, "error": str(e)}

    @staticmethod
    def parse_docx(path: str) -> Dict[str, Any]:
        doc = DocxDocument(path)
        paragraphs: List[Dict[str, Any]] = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            is_heading = style_name.lower().startswith("heading") or _is_heading(text)
            paragraphs.append({
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
            })
        
        tables: List[List[List[str]]] = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(rows)
        
        all_texts = [p["text"] for p in paragraphs]
        return {
            "type": "docx",
            "paragraphs": all_texts,
            "paragraph_meta": paragraphs,
            "tables": tables,
            "full_text": "\n".join(all_texts),
        }

    # ── Entity Extraction ─────────────────────
    @staticmethod
    def extract_entities(text: str) -> Dict[str, List[str]]:
        """Extract military entities from text with deduplication."""
        entities: Dict[str, List[str]] = {
            "officers": [],
            "ranks": [],
            "units": [],
            "operations": [],
            "weapons": [],
            "vehicles": [],
            "locations": [],
            "dates": [],
        }

        # Officers (rank + name)
        officers = _extract_officer_names(text)
        entities["officers"] = list(dict.fromkeys(officers))

        # Ranks (standalone)
        for pattern in RANK_PATTERNS:
            for m in re.finditer(pattern, text, re.IGNORECASE):
                rank = m.group(0).strip()
                if rank not in entities["ranks"]:
                    entities["ranks"].append(rank)

        # Units / formations
        for pattern in UNIT_PATTERNS:
            for m in re.finditer(pattern, text):
                unit = m.group(0).strip()
                if unit not in entities["units"]:
                    entities["units"].append(unit)

        # Operations / exercises
        for pattern in OPERATION_PATTERNS:
            for m in re.finditer(pattern, text):
                op = m.group(0).strip()
                if op not in entities["operations"]:
                    entities["operations"].append(op)

        # Weapons
        for pattern in WEAPON_PATTERNS:
            for m in re.finditer(pattern, text, re.IGNORECASE):
                w = m.group(0).strip()
                if w not in entities["weapons"]:
                    entities["weapons"].append(w)

        # Vehicles
        for pattern in VEHICLE_PATTERNS:
            for m in re.finditer(pattern, text, re.IGNORECASE):
                v = m.group(0).strip()
                if v not in entities["vehicles"]:
                    entities["vehicles"].append(v)

        # Locations
        for pattern in LOCATION_PATTERNS:
            for m in re.finditer(pattern, text):
                loc = m.group(0).strip()
                if loc not in entities["locations"]:
                    entities["locations"].append(loc)

        # Dates
        for pattern in DATE_PATTERNS:
            for m in re.finditer(pattern, text, re.IGNORECASE):
                d = m.group(0).strip()
                if d not in entities["dates"]:
                    entities["dates"].append(d)

        # Trim lists to avoid massive metadata payloads
        for key in entities:
            entities[key] = entities[key][:30]

        return entities

    # ── Smart Chunking ────────────────────────
    @staticmethod
    def to_chunks(
        parsed: Dict[str, Any],
        original_filename: str,
        chunk_size: int = 800,
        chunk_overlap: int = 150,
    ) -> List[Dict[str, Any]]:
        """
        Dispatch to the appropriate chunking strategy based on document type.
        All chunks carry rich metadata: source, page_number, section/heading,
        chunk_index, document_type, and extracted entity lists.
        """
        doc_type = parsed.get("type", "unknown")
        if doc_type in ("excel", "csv"):
            return DocumentProcessor._chunk_tabular(parsed, original_filename, chunk_size, chunk_overlap)
        elif doc_type == "pdf":
            return DocumentProcessor._chunk_pdf(parsed, original_filename, chunk_size, chunk_overlap)
        elif doc_type == "docx":
            return DocumentProcessor._chunk_docx(parsed, original_filename, chunk_size, chunk_overlap)
        return []

    # ── Tabular Chunking ──────────────────────
    @staticmethod
    def _chunk_tabular(
        parsed: Dict[str, Any],
        original_filename: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[Dict[str, Any]]:
        chunks: List[Dict[str, Any]] = []
        rows: List[str] = []

        if parsed["type"] == "excel":
            for sheet, records in parsed["sheets"].items():
                for record in records:
                    rows.append(f"Sheet: {sheet} | {json.dumps(record)}")
        else:
            for record in parsed["rows"]:
                rows.append(f"Row: {json.dumps(record)}")

        buf: List[str] = []
        buf_len = 0
        for row in rows:
            rlen = len(row)
            if buf_len + rlen > chunk_size and buf:
                content = "\n".join(buf)
                entities = DocumentProcessor.extract_entities(content)
                chunks.append({
                    "content": content,
                    "metadata": {
                        "source": original_filename,
                        "chunk_index": len(chunks),
                        "document_type": parsed["type"],
                        "entities": entities,
                    },
                })
                # Overlap: keep last N chars worth of rows
                overlap_buf: List[str] = []
                overlap_len = 0
                for prev_row in reversed(buf):
                    if overlap_len + len(prev_row) <= chunk_overlap:
                        overlap_buf.insert(0, prev_row)
                        overlap_len += len(prev_row)
                    else:
                        break
                buf = overlap_buf + [row]
                buf_len = overlap_len + rlen
            else:
                buf.append(row)
                buf_len += rlen

        if buf:
            content = "\n".join(buf)
            entities = DocumentProcessor.extract_entities(content)
            chunks.append({
                "content": content,
                "metadata": {
                    "source": original_filename,
                    "chunk_index": len(chunks),
                    "document_type": parsed["type"],
                    "entities": entities,
                },
            })

        return chunks

    # ── PDF Chunking (heading-aware + sentence-overlap) ──
    @staticmethod
    def _chunk_pdf(
        parsed: Dict[str, Any],
        original_filename: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[Dict[str, Any]]:
        """
        Chunk PDF with:
        - Per-page tracking for citation accuracy
        - Heading detection to start new semantic units
        - Sentence boundary overlap instead of character mid-cut
        - Table block preservation
        """
        chunks: List[Dict[str, Any]] = []
        current_section = "Document Start"

        for page_info in parsed.get("page_texts", []):
            page_num = page_info["page_number"]
            page_headings = page_info.get("headings", [])
            raw_text = page_info["text"]

            if not raw_text.strip():
                continue

            # Split into lines to detect headings and tables
            lines = raw_text.split("\n")
            sentences: List[str] = []
            i = 0

            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # Detect table block
                is_table, table_end = _detect_table_block(lines, i)
                if is_table and table_end > i:
                    table_text = "\n".join(lines[i:table_end]).strip()
                    if table_text:
                        sentences.append("[TABLE]\n" + table_text)
                    i = table_end
                    continue

                # Detect heading → update current section
                if _is_heading(line) or line in page_headings:
                    current_section = line
                    sentences.append(line)
                    i += 1
                    continue

                # Regular sentence accumulation (split on ". " boundaries)
                sub_sentences = re.split(r'(?<=[.!?])\s+', line)
                sentences.extend([s.strip() for s in sub_sentences if s.strip()])
                i += 1

            # Build chunks from sentences
            buf: List[str] = []
            buf_len = 0

            for sentence in sentences:
                slen = len(sentence)
                is_table_block = sentence.startswith("[TABLE]")

                # Table blocks always get their own chunk regardless of size
                if is_table_block:
                    if buf:
                        content = " ".join(buf)
                        entities = DocumentProcessor.extract_entities(content)
                        chunks.append({
                            "content": content,
                            "metadata": {
                                "source": original_filename,
                                "page_number": page_num,
                                "section": current_section,
                                "chunk_index": len(chunks),
                                "document_type": "pdf",
                                "chunk_type": "text",
                                "entities": entities,
                            },
                        })
                        buf = []
                        buf_len = 0
                    entities = DocumentProcessor.extract_entities(sentence)
                    chunks.append({
                        "content": sentence,
                        "metadata": {
                            "source": original_filename,
                            "page_number": page_num,
                            "section": current_section,
                            "chunk_index": len(chunks),
                            "document_type": "pdf",
                            "chunk_type": "table",
                            "entities": entities,
                        },
                    })
                    continue

                # Heading → flush buffer and start fresh
                if _is_heading(sentence):
                    if buf:
                        content = " ".join(buf)
                        entities = DocumentProcessor.extract_entities(content)
                        chunks.append({
                            "content": content,
                            "metadata": {
                                "source": original_filename,
                                "page_number": page_num,
                                "section": current_section,
                                "chunk_index": len(chunks),
                                "document_type": "pdf",
                                "chunk_type": "text",
                                "entities": entities,
                            },
                        })
                    current_section = sentence
                    buf = [sentence]
                    buf_len = slen
                    continue

                # Normal overflow → flush with sentence-boundary overlap
                if buf_len + slen > chunk_size and buf:
                    content = " ".join(buf)
                    entities = DocumentProcessor.extract_entities(content)
                    chunks.append({
                        "content": content,
                        "metadata": {
                            "source": original_filename,
                            "page_number": page_num,
                            "section": current_section,
                            "chunk_index": len(chunks),
                            "document_type": "pdf",
                            "chunk_type": "text",
                            "entities": entities,
                        },
                    })
                    # Sentence-boundary overlap
                    overlap_buf: List[str] = []
                    overlap_len = 0
                    for prev in reversed(buf):
                        if overlap_len + len(prev) <= chunk_overlap:
                            overlap_buf.insert(0, prev)
                            overlap_len += len(prev)
                        else:
                            break
                    buf = overlap_buf + [sentence]
                    buf_len = overlap_len + slen
                else:
                    buf.append(sentence)
                    buf_len += slen

            # Flush remaining buffer for this page
            if buf:
                content = " ".join(buf)
                entities = DocumentProcessor.extract_entities(content)
                chunks.append({
                    "content": content,
                    "metadata": {
                        "source": original_filename,
                        "page_number": page_num,
                        "section": current_section,
                        "chunk_index": len(chunks),
                        "document_type": "pdf",
                        "chunk_type": "text",
                        "entities": entities,
                    },
                })

        return chunks

    # ── DOCX Chunking (heading-aware + paragraph-aware) ──
    @staticmethod
    def _chunk_docx(
        parsed: Dict[str, Any],
        original_filename: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[Dict[str, Any]]:
        chunks: List[Dict[str, Any]] = []
        current_section = "Document Start"
        buf: List[str] = []
        buf_len = 0

        para_meta = parsed.get("paragraph_meta", [])
        if not para_meta:
            # Fallback: treat plain paragraphs list
            para_meta = [{"text": p, "is_heading": _is_heading(p), "style": ""} for p in parsed.get("paragraphs", [])]

        # Process tables first: serialize them as special chunks
        table_chunks: List[Dict[str, Any]] = []
        for table_idx, table in enumerate(parsed.get("tables", [])):
            table_text = f"[TABLE {table_idx + 1}]\n"
            for row in table:
                table_text += " | ".join(str(c) for c in row) + "\n"
            entities = DocumentProcessor.extract_entities(table_text)
            table_chunks.append({
                "content": table_text.strip(),
                "metadata": {
                    "source": original_filename,
                    "section": current_section,
                    "chunk_index": -1,
                    "document_type": "docx",
                    "chunk_type": "table",
                    "entities": entities,
                },
            })

        def flush_buf():
            nonlocal buf, buf_len
            if buf:
                content = "\n".join(buf)
                entities = DocumentProcessor.extract_entities(content)
                chunks.append({
                    "content": content,
                    "metadata": {
                        "source": original_filename,
                        "section": current_section,
                        "chunk_index": len(chunks),
                        "document_type": "docx",
                        "chunk_type": "text",
                        "entities": entities,
                    },
                })
                buf = []
                buf_len = 0

        for pm in para_meta:
            text = pm["text"]
            is_heading = pm.get("is_heading", False)
            tlen = len(text)

            if is_heading:
                flush_buf()
                current_section = text
                buf = [text]
                buf_len = tlen
                continue

            if buf_len + tlen > chunk_size and buf:
                flush_buf()
                # Sentence-boundary overlap
                overlap_buf: List[str] = []
                overlap_len = 0
                for prev in reversed(buf):
                    if overlap_len + len(prev) <= chunk_overlap:
                        overlap_buf.insert(0, prev)
                        overlap_len += len(prev)
                    else:
                        break
                buf = overlap_buf + [text]
                buf_len = overlap_len + tlen
            else:
                buf.append(text)
                buf_len += tlen

        flush_buf()

        # Re-index table chunks and append
        for tc in table_chunks:
            tc["metadata"]["chunk_index"] = len(chunks)
            chunks.append(tc)

        return chunks